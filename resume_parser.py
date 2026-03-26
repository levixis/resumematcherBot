"""
Resume Parser Module
Handles extracting text from PDF, DOCX, and TXT files.
"""

import os
from PyPDF2 import PdfReader
from docx import Document


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")


def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")


def parse_txt(file_path: str) -> str:
    """Extract text from a TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as e:
        raise ValueError(f"Failed to parse TXT: {e}")


def parse_resume(file_path: str) -> str:
    """
    Auto-detect file format and extract text.
    Supports: PDF, DOCX, TXT
    """
    ext = os.path.splitext(file_path)[1].lower()

    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".txt": parse_txt,
        ".text": parse_txt,
        ".doc": parse_txt,  # Fallback for .doc (limited support)
    }

    parser = parsers.get(ext)
    if parser is None:
        # Try to read as plain text as fallback
        try:
            return parse_txt(file_path)
        except Exception:
            raise ValueError(
                f"Unsupported file format: {ext}\n"
                "Supported formats: PDF, DOCX, TXT"
            )

    text = parser(file_path)
    if not text:
        raise ValueError(
            "Could not extract any text from the file. "
            "Please ensure the file is not empty or image-based."
        )
    return text
