"""
CV Generator Module
Generates professionally formatted CVs in multiple TEMPLATE STYLES and export formats.

Template Styles:
  1. Classic      — Traditional single-column, formal layout
  2. Modern       — Two-column layout with sidebar for skills/contact
  3. Minimal      — Clean, whitespace-heavy, minimalist design
  4. ATS-Friendly — Plain, keyword-optimized, guaranteed to pass ATS scanners
"""

import os
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib.colors import HexColor, Color
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Frame, PageTemplate, BaseDocTemplate
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ═══════════════════════════════════════════════════════════════
#  AVAILABLE TEMPLATES
# ═══════════════════════════════════════════════════════════════
TEMPLATES = {
    "classic": {
        "name": "📜 Classic",
        "description": "Traditional single-column formal layout",
        "primary": "#1a365d",
        "accent": "#2b6cb0",
        "text": "#2d3748",
        "light": "#e2e8f0",
    },
    "modern": {
        "name": "🎨 Modern Two-Column",
        "description": "Sleek two-column layout with colored sidebar",
        "primary": "#0f172a",
        "accent": "#6366f1",
        "text": "#334155",
        "light": "#e0e7ff",
        "sidebar_bg": "#1e293b",
        "sidebar_text": "#e2e8f0",
    },
    "minimal": {
        "name": "✨ Minimal",
        "description": "Clean, whitespace-heavy minimalist design",
        "primary": "#111827",
        "accent": "#6b7280",
        "text": "#374151",
        "light": "#f3f4f6",
    },
    "ats": {
        "name": "🤖 ATS-Friendly",
        "description": "Plain text optimized for Applicant Tracking Systems",
        "primary": "#000000",
        "accent": "#333333",
        "text": "#000000",
        "light": "#cccccc",
    },
}


def get_template_list() -> list:
    """Return list of available templates for display."""
    return [
        {"id": k, "name": v["name"], "description": v["description"]}
        for k, v in TEMPLATES.items()
    ]


# ═══════════════════════════════════════════════════════════════
#  PDF GENERATORS
# ═══════════════════════════════════════════════════════════════

def _make_styles(template_id: str):
    """Build ReportLab styles for a given template."""
    t = TEMPLATES[template_id]
    primary = HexColor(t["primary"])
    accent = HexColor(t["accent"])
    text_color = HexColor(t["text"])
    light = HexColor(t["light"])

    styles = getSampleStyleSheet()

    # Name sizes vary by template
    name_size = 22 if template_id != "minimal" else 28
    body_size = 10 if template_id != "ats" else 11

    styles.add(ParagraphStyle(
        name="CVName", fontSize=name_size, leading=name_size + 4,
        textColor=primary, fontName="Helvetica-Bold",
        alignment=TA_LEFT, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="CVSection", fontSize=12, leading=16,
        textColor=primary, fontName="Helvetica-Bold",
        spaceBefore=14, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="CVBody", fontSize=body_size, leading=body_size + 4,
        textColor=text_color, fontName="Helvetica",
        alignment=TA_JUSTIFY, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="CVBullet", fontSize=body_size, leading=body_size + 3,
        textColor=text_color, fontName="Helvetica",
        leftIndent=15, spaceAfter=2, bulletIndent=5,
    ))
    styles.add(ParagraphStyle(
        name="CVSub", fontSize=11, leading=14,
        textColor=HexColor("#1a202c"), fontName="Helvetica-Bold",
        spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        name="CVMeta", fontSize=9, leading=12,
        textColor=accent, fontName="Helvetica-Oblique",
        spaceAfter=4,
    ))
    # Sidebar styles (for modern template)
    styles.add(ParagraphStyle(
        name="SidebarTitle", fontSize=11, leading=14,
        textColor=HexColor("#e2e8f0"), fontName="Helvetica-Bold",
        spaceBefore=10, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="SidebarBody", fontSize=9, leading=12,
        textColor=HexColor("#cbd5e1"), fontName="Helvetica",
        spaceAfter=2,
    ))

    return styles, primary, accent, text_color, light


def _add_section_pdf(story, title, styles, primary, light, template_id):
    """Add a section heading with appropriate divider for the template."""
    if template_id == "ats":
        story.append(Spacer(1, 8))
        story.append(Paragraph(title.upper(), styles["CVSection"]))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#000000"),
                                spaceBefore=0, spaceAfter=6))
    elif template_id == "minimal":
        story.append(Spacer(1, 16))
        story.append(Paragraph(title.upper(), styles["CVSection"]))
        story.append(Spacer(1, 4))
    else:
        story.append(HRFlowable(width="100%", thickness=0.5, color=light,
                                spaceBefore=10, spaceAfter=4))
        story.append(Paragraph(f"<b>{title.upper()}</b>", styles["CVSection"]))


def generate_classic_pdf(content, path, name):
    """Classic single-column PDF."""
    doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=20*mm, leftMargin=20*mm,
                            topMargin=15*mm, bottomMargin=15*mm)
    styles, primary, accent, text_color, light = _make_styles("classic")
    story = []

    story.append(Paragraph(name.upper(), styles["CVName"]))
    story.append(HRFlowable(width="100%", thickness=2, color=primary,
                            spaceBefore=4, spaceAfter=10))

    _add_common_sections(story, content, styles, primary, light, "classic")
    doc.build(story)
    return path


def generate_modern_pdf(content, path, name):
    """Modern two-column PDF with sidebar."""
    doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=10*mm, leftMargin=10*mm,
                            topMargin=10*mm, bottomMargin=10*mm)
    styles, primary, accent, text_color, light = _make_styles("modern")

    # Build sidebar content
    sidebar_content = []
    sidebar_content.append(Paragraph(name.upper(), ParagraphStyle(
        name="SBName", fontSize=16, leading=20, textColor=HexColor("#ffffff"),
        fontName="Helvetica-Bold", spaceAfter=10,
    )))

    # Skills in sidebar
    skills = content.get("skills", [])
    if skills:
        sidebar_content.append(Paragraph("SKILLS", styles["SidebarTitle"]))
        for skill in skills:
            sidebar_content.append(Paragraph(f"▸ {skill}", styles["SidebarBody"]))

    # Certifications in sidebar
    certs = content.get("certifications", [])
    if certs:
        sidebar_content.append(Spacer(1, 10))
        sidebar_content.append(Paragraph("CERTIFICATIONS", styles["SidebarTitle"]))
        for cert in certs:
            sidebar_content.append(Paragraph(f"▸ {cert}", styles["SidebarBody"]))

    # Education in sidebar
    education = content.get("education", [])
    if education:
        sidebar_content.append(Spacer(1, 10))
        sidebar_content.append(Paragraph("EDUCATION", styles["SidebarTitle"]))
        for edu in education:
            sidebar_content.append(Paragraph(
                f"<b>{edu.get('degree', '')}</b>", styles["SidebarBody"]))
            sidebar_content.append(Paragraph(
                edu.get("institution", ""), styles["SidebarBody"]))
            sidebar_content.append(Paragraph(
                edu.get("year", ""), styles["SidebarBody"]))
            sidebar_content.append(Spacer(1, 4))

    # Build main content
    main_content = []

    summary = content.get("professional_summary", "")
    if summary:
        main_content.append(Paragraph("<b>PROFESSIONAL SUMMARY</b>", styles["CVSection"]))
        main_content.append(Paragraph(summary, styles["CVBody"]))

    experience = content.get("experience", [])
    if experience:
        main_content.append(Spacer(1, 8))
        main_content.append(Paragraph("<b>EXPERIENCE</b>", styles["CVSection"]))
        for exp in experience:
            main_content.append(Paragraph(exp.get("title", ""), styles["CVSub"]))
            main_content.append(Paragraph(
                f"{exp.get('company', '')}  |  {exp.get('duration', '')}",
                styles["CVMeta"]))
            for bullet in exp.get("bullets", []):
                main_content.append(Paragraph(f"▸  {bullet}", styles["CVBullet"]))
            main_content.append(Spacer(1, 4))

    projects = content.get("projects", [])
    if projects:
        main_content.append(Spacer(1, 8))
        main_content.append(Paragraph("<b>KEY PROJECTS</b>", styles["CVSection"]))
        for proj in projects:
            main_content.append(Paragraph(
                f"<b>{proj.get('name', '')}</b>", styles["CVSub"]))
            main_content.append(Paragraph(proj.get("description", ""), styles["CVBody"]))

    # Create two-column table
    sidebar_bg = HexColor(TEMPLATES["modern"]["sidebar_bg"])
    page_w = A4[0] - 20*mm
    sidebar_w = page_w * 0.32
    main_w = page_w * 0.65

    table_data = [[sidebar_content, main_content]]
    table = Table(table_data, colWidths=[sidebar_w, main_w])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), sidebar_bg),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (0, 0), 12),
        ("RIGHTPADDING", (0, 0), (0, 0), 10),
        ("LEFTPADDING", (1, 0), (1, 0), 14),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
        ("ROUNDEDCORNERS", [6, 6, 6, 6]),
    ]))

    doc.build([table])
    return path


def generate_minimal_pdf(content, path, name):
    """Minimal clean PDF."""
    doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=30*mm, leftMargin=30*mm,
                            topMargin=25*mm, bottomMargin=20*mm)
    styles, primary, accent, text_color, light = _make_styles("minimal")
    story = []

    story.append(Paragraph(name, styles["CVName"]))
    story.append(Spacer(1, 20))

    _add_common_sections(story, content, styles, primary, light, "minimal")
    doc.build(story)
    return path


def generate_ats_pdf(content, path, name):
    """ATS-friendly plain PDF."""
    doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=20*mm, leftMargin=20*mm,
                            topMargin=15*mm, bottomMargin=15*mm)
    styles, primary, accent, text_color, light = _make_styles("ats")
    story = []

    story.append(Paragraph(name.upper(), styles["CVName"]))
    _add_common_sections(story, content, styles, primary, light, "ats")
    doc.build(story)
    return path


def _add_common_sections(story, content, styles, primary, light, template_id):
    """Add common resume sections to the PDF story."""
    summary = content.get("professional_summary", "")
    if summary:
        _add_section_pdf(story, "Professional Summary", styles, primary, light, template_id)
        story.append(Paragraph(summary, styles["CVBody"]))

    skills = content.get("skills", [])
    if skills:
        _add_section_pdf(story, "Skills", styles, primary, light, template_id)
        joiner = "  |  " if template_id == "ats" else "  •  "
        story.append(Paragraph(joiner.join(skills), styles["CVBody"]))

    experience = content.get("experience", [])
    if experience:
        _add_section_pdf(story, "Experience", styles, primary, light, template_id)
        for exp in experience:
            story.append(Paragraph(exp.get("title", ""), styles["CVSub"]))
            story.append(Paragraph(
                f"{exp.get('company', '')}  |  {exp.get('duration', '')}",
                styles["CVMeta"]))
            for bullet in exp.get("bullets", []):
                marker = "-" if template_id == "ats" else "▸"
                story.append(Paragraph(f"{marker}  {bullet}", styles["CVBullet"]))
            story.append(Spacer(1, 4))

    projects = content.get("projects", [])
    if projects:
        _add_section_pdf(story, "Projects", styles, primary, light, template_id)
        for proj in projects:
            story.append(Paragraph(
                f"<b>{proj.get('name', '')}</b>", styles["CVSub"]))
            story.append(Paragraph(proj.get("description", ""), styles["CVBody"]))
            story.append(Spacer(1, 3))

    education = content.get("education", [])
    if education:
        _add_section_pdf(story, "Education", styles, primary, light, template_id)
        for edu in education:
            story.append(Paragraph(
                f"<b>{edu.get('degree', '')}</b>", styles["CVSub"]))
            story.append(Paragraph(
                f"{edu.get('institution', '')}  |  {edu.get('year', '')}",
                styles["CVMeta"]))

    certs = content.get("certifications", [])
    if certs:
        _add_section_pdf(story, "Certifications", styles, primary, light, template_id)
        for cert in certs:
            marker = "-" if template_id == "ats" else "▸"
            story.append(Paragraph(f"{marker}  {cert}", styles["CVBullet"]))


# ═══════════════════════════════════════════════════════════════
#  DOCX GENERATORS
# ═══════════════════════════════════════════════════════════════

def _docx_add_section(doc, title, template_id):
    """Add a styled section heading in DOCX."""
    t = TEMPLATES[template_id]
    pr = t["primary"]
    r, g, b = int(pr[1:3], 16), int(pr[3:5], 16), int(pr[5:7], 16)

    if template_id != "minimal":
        line_para = doc.add_paragraph()
        line_para.space_before = Pt(10)
        line_para.space_after = Pt(0)
        line_run = line_para.add_run("─" * 80)
        line_run.font.size = Pt(6)
        line_run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    heading_para = doc.add_paragraph()
    heading_para.space_before = Pt(6 if template_id != "minimal" else 14)
    heading_para.space_after = Pt(6)
    heading_run = heading_para.add_run(title.upper())
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    heading_run.font.color.rgb = RGBColor(r, g, b)


def _docx_body(doc, text, template_id):
    t = TEMPLATES[template_id]
    tc = t["text"]
    r, g, b = int(tc[1:3], 16), int(tc[3:5], 16), int(tc[5:7], 16)
    para = doc.add_paragraph()
    para.space_after = Pt(4)
    run = para.add_run(text)
    run.font.size = Pt(10 if template_id != "ats" else 11)
    run.font.color.rgb = RGBColor(r, g, b)


def _docx_subheading(doc, text, template_id):
    para = doc.add_paragraph()
    para.space_before = Pt(4)
    para.space_after = Pt(1)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)


def _docx_meta(doc, text, template_id):
    t = TEMPLATES[template_id]
    ac = t["accent"]
    r, g, b = int(ac[1:3], 16), int(ac[3:5], 16), int(ac[5:7], 16)
    para = doc.add_paragraph()
    para.space_after = Pt(3)
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(r, g, b)


def _docx_bullet(doc, text, template_id):
    marker = "-" if template_id == "ats" else "▸"
    para = doc.add_paragraph()
    para.space_after = Pt(2)
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(f"{marker}  {text}")
    run.font.size = Pt(10 if template_id != "ats" else 11)
    run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)


def _generate_standard_docx(content, path, name, template_id):
    """Generate DOCX for classic / minimal / ats templates."""
    doc = Document()
    t = TEMPLATES[template_id]
    pr = t["primary"]
    r, g, b = int(pr[1:3], 16), int(pr[3:5], 16), int(pr[5:7], 16)

    margin = 0.8 if template_id != "minimal" else 1.0
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    display_name = name.upper() if template_id != "minimal" else name
    name_run = name_para.add_run(display_name)
    name_run.bold = True
    name_run.font.size = Pt(22 if template_id != "minimal" else 28)
    name_run.font.color.rgb = RGBColor(r, g, b)
    name_para.space_after = Pt(4)

    # Summary
    summary = content.get("professional_summary", "")
    if summary:
        _docx_add_section(doc, "Professional Summary", template_id)
        _docx_body(doc, summary, template_id)

    # Skills
    skills = content.get("skills", [])
    if skills:
        _docx_add_section(doc, "Skills", template_id)
        joiner = "  |  " if template_id == "ats" else "  •  "
        _docx_body(doc, joiner.join(skills), template_id)

    # Experience
    experience = content.get("experience", [])
    if experience:
        _docx_add_section(doc, "Experience", template_id)
        for exp in experience:
            _docx_subheading(doc, exp.get("title", ""), template_id)
            _docx_meta(doc, f"{exp.get('company', '')}  |  {exp.get('duration', '')}", template_id)
            for bullet in exp.get("bullets", []):
                _docx_bullet(doc, bullet, template_id)

    # Projects
    projects = content.get("projects", [])
    if projects:
        _docx_add_section(doc, "Projects", template_id)
        for proj in projects:
            _docx_subheading(doc, proj.get("name", ""), template_id)
            _docx_body(doc, proj.get("description", ""), template_id)

    # Education
    education = content.get("education", [])
    if education:
        _docx_add_section(doc, "Education", template_id)
        for edu in education:
            _docx_subheading(doc, edu.get("degree", ""), template_id)
            _docx_meta(doc, f"{edu.get('institution', '')}  |  {edu.get('year', '')}", template_id)

    # Certifications
    certs = content.get("certifications", [])
    if certs:
        _docx_add_section(doc, "Certifications", template_id)
        for cert in certs:
            _docx_bullet(doc, cert, template_id)

    doc.save(path)
    return path


def _generate_modern_docx(content, path, name):
    """Generate a two-column DOCX using a table layout."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    # Create a 1-row, 2-column table for the layout
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for cell in table.columns[0].cells:
        cell.width = Inches(2.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(4.8)

    # ── LEFT SIDEBAR ──
    left_cell = table.cell(0, 0)

    # Set sidebar background color
    shading = left_cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): '1e293b'
    })
    shading.append(shading_elem)

    # Name in sidebar
    p = left_cell.paragraphs[0]
    p.space_after = Pt(10)
    run = p.add_run(name.upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def sidebar_heading(text):
        p = left_cell.add_paragraph()
        p.space_before = Pt(12)
        p.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)

    def sidebar_item(text):
        p = left_cell.add_paragraph()
        p.space_after = Pt(2)
        run = p.add_run(f"▸ {text}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

    # Skills
    skills = content.get("skills", [])
    if skills:
        sidebar_heading("SKILLS")
        for skill in skills:
            sidebar_item(skill)

    # Education
    education = content.get("education", [])
    if education:
        sidebar_heading("EDUCATION")
        for edu in education:
            p = left_cell.add_paragraph()
            p.space_after = Pt(1)
            run = p.add_run(edu.get("degree", ""))
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
            p2 = left_cell.add_paragraph()
            p2.space_after = Pt(1)
            run2 = p2.add_run(edu.get("institution", ""))
            run2.font.size = Pt(8)
            run2.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            p3 = left_cell.add_paragraph()
            p3.space_after = Pt(6)
            run3 = p3.add_run(edu.get("year", ""))
            run3.font.size = Pt(8)
            run3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # Certifications
    certs = content.get("certifications", [])
    if certs:
        sidebar_heading("CERTIFICATIONS")
        for cert in certs:
            sidebar_item(cert)

    # ── RIGHT MAIN CONTENT ──
    right_cell = table.cell(0, 1)

    def main_heading(text):
        p = right_cell.add_paragraph()
        p.space_before = Pt(10)
        p.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    def main_body(text):
        p = right_cell.add_paragraph()
        p.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    def main_sub(text):
        p = right_cell.add_paragraph()
        p.space_before = Pt(4)
        p.space_after = Pt(1)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)

    def main_meta(text):
        p = right_cell.add_paragraph()
        p.space_after = Pt(3)
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)

    def main_bullet(text):
        p = right_cell.add_paragraph()
        p.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(f"▸  {text}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Remove default empty paragraph
    if right_cell.paragraphs[0].text == "":
        right_cell.paragraphs[0].clear()

    summary = content.get("professional_summary", "")
    if summary:
        main_heading("PROFESSIONAL SUMMARY")
        main_body(summary)

    experience = content.get("experience", [])
    if experience:
        main_heading("EXPERIENCE")
        for exp in experience:
            main_sub(exp.get("title", ""))
            main_meta(f"{exp.get('company', '')}  |  {exp.get('duration', '')}")
            for bullet in exp.get("bullets", []):
                main_bullet(bullet)

    projects = content.get("projects", [])
    if projects:
        main_heading("KEY PROJECTS")
        for proj in projects:
            main_sub(proj.get("name", ""))
            main_body(proj.get("description", ""))

    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders.append(border)
    tblPr.append(borders)

    doc.save(path)
    return path


# ═══════════════════════════════════════════════════════════════
#  PUBLIC API
# ═══════════════════════════════════════════════════════════════

def generate_pdf(content: dict, output_path: str, candidate_name: str = "Candidate",
                 template_id: str = "classic") -> str:
    """
    Generate a PDF CV in the specified template style.
    
    Args:
        content: Optimized resume content dict from matcher
        output_path: Where to save the PDF
        candidate_name: Name for the CV header
        template_id: One of 'classic', 'modern', 'minimal', 'ats'
    """
    generators = {
        "classic": generate_classic_pdf,
        "modern": generate_modern_pdf,
        "minimal": generate_minimal_pdf,
        "ats": generate_ats_pdf,
    }
    gen = generators.get(template_id, generate_classic_pdf)
    return gen(content, output_path, candidate_name)


def generate_docx(content: dict, output_path: str, candidate_name: str = "Candidate",
                  template_id: str = "classic") -> str:
    """
    Generate a DOCX CV in the specified template style.
    
    Args:
        content: Optimized resume content dict from matcher
        output_path: Where to save the DOCX
        candidate_name: Name for the CV header
        template_id: One of 'classic', 'modern', 'minimal', 'ats'
    """
    if template_id == "modern":
        return _generate_modern_docx(content, output_path, candidate_name)
    else:
        return _generate_standard_docx(content, output_path, candidate_name, template_id)
